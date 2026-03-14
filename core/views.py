from django.shortcuts import render, redirect
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.http import JsonResponse
import json


from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Avg

# Modellarni import qilish
from elementar.models import Lesson as ELesson, Element as EElement, UserProgress as EProgress
from adaptive.models import Course as ACourse, Element as AElement, Progress as AProgress
from neural.models import AIPath as NPath, Element as NElement, NeuralProgress as NProgress
from quiz.models import ExamSession


from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from django.db.models import Avg, Max

# Modellarni import qilish
from elementar.models import Element as EElement, UserProgress as EProgress
from adaptive.models import Element as AElement, Progress as AProgress
from neural.models import Element as NElement, NeuralProgress as NProgress
from quiz.models import ExamSession
from practical.models import Submission as PSubmission

@login_required
def dashboard(request):
    user = request.user

    # --- 1. ELEMENTAR, ADAPTIVE, NEURAL PROGRESS ---
    e_total = EElement.objects.count()
    e_done = EProgress.objects.filter(user=user, is_completed=True).count()
    e_percent = int((e_done / e_total) * 100) if e_total > 0 else 0

    a_total = AElement.objects.count()
    a_done = AProgress.objects.filter(user=user, completed=True).count()
    a_percent = int((a_done / a_total) * 100) if a_total > 0 else 0

    n_total = NElement.objects.filter(is_essential=True).count()
    n_done = NProgress.objects.filter(user=user, completed=True, element__is_essential=True).count()
    n_percent = int((n_done / n_total) * 100) if n_total > 0 else 0

    # --- 2. NAZORAT BALLARI MANTIQI (SIYaSAT) ---
    
    # Oraliq nazorat: E, A, N progresslari ichida eng yuqorisi
    oraliq_ball = float(max(e_percent, a_percent, n_percent))

    # Joriy nazorat: Practical app dagi baholangan topshiriqlar o'rtachasi
    practical_graded = PSubmission.objects.filter(user=user, is_graded=True)
    joriy_ball = float(practical_graded.aggregate(Avg('grade'))['grade__avg'] or 0)
    
    # Yakuniy nazorat: Quiz app dagi test natijalari o'rtachasi
    quiz_sessions = ExamSession.objects.filter(user=user, is_finished=True)
    yakuniy_ball = float(quiz_sessions.aggregate(Avg('score'))['score__avg'] or 0)

    # UMUMIY NATIJA FORMULASI:
    # ((Oraliq + Joriy) / 2) ballarning 50% ini beradi
    # Yakuniy ball qolgan 50% ini beradi
    total_mastery = int(((oraliq_ball + joriy_ball) / 2 * 0.5) + (yakuniy_ball * 0.5))

    # Detallar
    recent_practicals = practical_graded.order_by('-submitted_at')[:3]
    recent_quizzes = quiz_sessions.order_by('-started_at')[:3]

    context = {
        'e_percent': e_percent, 'a_percent': a_percent, 'n_percent': n_percent,
        'oraliq_ball': int(oraliq_ball),
        'joriy_ball': int(joriy_ball),
        'yakuniy_ball': int(yakuniy_ball),
        'total_mastery': total_mastery,
        'e_done': e_done, 'e_total': e_total,
        'a_done': a_done, 'a_total': a_total,
        'n_done': n_done, 'n_total': n_total,
        'recent_practicals': recent_practicals,
        'recent_quizzes': recent_quizzes,
    }
    return render(request, 'core/dashboard.html', context)


# Bosh sahifa
def home(request):
    return render(request, 'core/home.html')

# Ro'yxatdan o'tish
def register_view(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('core:home')
    else:
        form = UserCreationForm()
    return render(request, 'core/register.html', {'form': form})

# Tizimga kirish
def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('core:home')
    else:
        form = AuthenticationForm()
    return render(request, 'core/login.html', {'form': form})

# Tizimdan chiqish
def logout_view(request):
    logout(request)
    return redirect('core:home')


import json
from django.http import JsonResponse
from .models import BotKnowledge

def chatbot_api(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            user_message = data.get('message', '').lower()
            
            # Bazadagi barcha bilimlarni olamiz
            # (Agar baza juda katta bo'lsa, buni optimallashtirish mumkin)
            knowledges = BotKnowledge.objects.all()
            
            reply = None
            
            # Foydalanuvchi xabarida qatnashgan so'zlarga qarab bazadan qidiramiz
            for k in knowledges:
                # Agar bazadagi kalit so'z user xabarining ichida bo'lsa
                if k.keyword.lower() in user_message:
                    reply = k.answer
                    break # Birinchi topilgan moslikni olamiz
            
            # Agar mos keladigan hech narsa topilmasa
            if not reply:
                reply = "Kechirasiz, bu termin bo'yicha menda ma'lumot yo'q. Iltimos, boshqacharoq so'rab ko'ring yoki administratorga murojaat qiling."
                
            return JsonResponse({'reply': reply})
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)
            
    return JsonResponse({'error': 'Faqat POST so\'rovlar qabul qilinadi'}, status=400)



import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
from django.utils import timezone

# ... oldingi importlar ...

@login_required
def export_results_docx(request):
    user = request.user

    # --- 1. MA'LUMOTLARNI QAYTADAN HISOBLASH ---
    # Elementar
    e_total = EElement.objects.count()
    e_done = EProgress.objects.filter(user=user, is_completed=True).count()
    e_percent = int((e_done / e_total) * 100) if e_total > 0 else 0

    # Adaptive
    a_total = AElement.objects.count()
    a_done = AProgress.objects.filter(user=user, completed=True).count()
    a_percent = int((a_done / a_total) * 100) if a_total > 0 else 0

    # Neural
    n_total = NElement.objects.filter(is_essential=True).count()
    n_done = NProgress.objects.filter(user=user, completed=True, element__is_essential=True).count()
    n_percent = int((n_done / n_total) * 100) if n_total > 0 else 0

    # Nazorat ballari
    oraliq_ball = max(e_percent, a_percent, n_percent)
    
    practical_graded = PSubmission.objects.filter(user=user, is_graded=True)
    joriy_ball = practical_graded.aggregate(Avg('grade'))['grade__avg'] or 0
    
    quiz_sessions = ExamSession.objects.filter(user=user, is_finished=True)
    yakuniy_ball = quiz_sessions.aggregate(Avg('score'))['score__avg'] or 0

    total_mastery = int((oraliq_ball + joriy_ball + yakuniy_ball) / 3)

    # --- 2. WORD DOKUMENTNI YARATISH ---
    document = Document()
    
    header = document.add_heading('ILMS Platformasi - O\'quv Natijalari', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph()
    p.add_run(f"Talaba: ").bold = True
    p.add_run(f"{user.get_full_name() or user.username}")
    p.add_run(f"\nSana: ").bold = True
    p.add_run(f"{timezone.now().strftime('%d.%m.%Y %H:%M')}")

    document.add_heading('Nazorat ko\'rsatkichlari jadvali', level=1)
    
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nazorat turi'
    hdr_cells[1].text = 'Natija'

    # Ma'lumotlarni jadvalga qo'shish
    data = [
        ('Joriy Nazorat (Amaliy topshiriqlar)', f"{int(joriy_ball)} ball"),
        ('Oraliq Nazorat (O\'quv modullari)', f"{int(oraliq_ball)} ball"),
        ('Yakuniy Nazorat (Test sinovlari)', f"{int(yakuniy_ball)} ball"),
        ('UMUMIY O\'ZLASHTIRISH DARAJASI', f"{total_mastery}%"),
    ]

    for label, value in data:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    # Faylni yuborish
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Natijalar_{user.username}.docx'
    return response



@login_required
def view_certificate(request):
    user = request.user

    # --- 1. MA'LUMOTLARNI QAYTADAN HISOBLASH (Xavfsizlik uchun) ---
    # Elementar
    e_total = EElement.objects.count()
    e_done = EProgress.objects.filter(user=user, is_completed=True).count()
    e_percent = int((e_done / e_total) * 100) if e_total > 0 else 0

    # Adaptive
    a_total = AElement.objects.count()
    a_done = AProgress.objects.filter(user=user, completed=True).count()
    a_percent = int((a_done / a_total) * 100) if a_total > 0 else 0

    # Neural
    n_total = NElement.objects.filter(is_essential=True).count()
    n_done = NProgress.objects.filter(user=user, completed=True, element__is_essential=True).count()
    n_percent = int((n_done / n_total) * 100) if n_total > 0 else 0

    # Nazorat ballari mantiqi
    oraliq_ball = max(e_percent, a_percent, n_percent)
    
    practical_graded = PSubmission.objects.filter(user=user, is_graded=True)
    joriy_ball = practical_graded.aggregate(Avg('grade'))['grade__avg'] or 0
    
    quiz_sessions = ExamSession.objects.filter(user=user, is_finished=True)
    yakuniy_ball = quiz_sessions.aggregate(Avg('score'))['score__avg'] or 0

    # Umumiy natija
    total_mastery = int((oraliq_ball + joriy_ball + yakuniy_ball) / 3)

    # --- 2. TEKSHIRUV VA RENDER ---
    if total_mastery < 60:
        return HttpResponse("Kechirasiz, sertifikat olish uchun natijangiz kamida 60% bo'lishi kerak.", status=403)

    return render(request, 'core/certificate.html', {
        'total_mastery': total_mastery,
        'date': timezone.now(),
    })